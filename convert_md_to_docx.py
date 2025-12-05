#!/usr/bin/env python3
"""
Convert Markdown to Word Document (.docx)
"""
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


def parse_markdown_table(lines):
    """Parse markdown table lines into rows."""
    rows = []
    for line in lines:
        if '|' in line and not line.strip().startswith('|--') and not re.match(r'^\|[\s\-:]+\|$', line.strip()):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)
    return rows


def add_table_to_doc(doc, rows):
    """Add a table to the document."""
    if not rows:
        return
    
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = cell_text
                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    doc.add_paragraph()


def convert_md_to_docx(md_path, docx_path):
    """Convert markdown file to Word document."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set document title
    title = doc.add_heading('PRÉSENTATION OFFICIELLE DE LA PLATEFORME TAXCOLLECTOR', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lines = content.split('\n')
    i = 0
    table_lines = []
    in_table = False
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip the first title (already added)
        if stripped.startswith('# PRÉSENTATION OFFICIELLE'):
            i += 1
            continue
        
        # Handle tables
        if '|' in stripped and stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(stripped)
            i += 1
            continue
        elif in_table:
            # End of table
            rows = parse_markdown_table(table_lines)
            add_table_to_doc(doc, rows)
            in_table = False
            table_lines = []
        
        # Handle headings
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=3)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        
        # Handle horizontal rules
        elif stripped == '---':
            doc.add_paragraph('─' * 50)
        
        # Handle bullet points
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            # Handle bold text
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
        
        # Handle regular paragraphs
        elif stripped and not stripped.startswith('```'):
            # Clean up markdown formatting
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', stripped)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            text = re.sub(r'`([^`]+)`', r'\1', text)
            
            if text:
                p = doc.add_paragraph(text)
        
        # Handle code blocks (skip them or add as plain text)
        elif stripped.startswith('```'):
            # Skip code block markers
            pass
        
        i += 1
    
    # Handle any remaining table
    if in_table and table_lines:
        rows = parse_markdown_table(table_lines)
        add_table_to_doc(doc, rows)
    
    # Save document
    doc.save(docx_path)
    print(f"Document saved to: {docx_path}")


if __name__ == '__main__':
    md_path = '.trae/documents/PRESENTATION_OFFICIELLE_TAXCOLLECTOR.md'
    docx_path = '.trae/documents/PRESENTATION_OFFICIELLE_TAXCOLLECTOR.docx'
    convert_md_to_docx(md_path, docx_path)
